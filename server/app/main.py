import os
import base64
import re
import shutil
import tempfile
import time
import uuid
from datetime import datetime
from contextlib import asynccontextmanager
from io import BytesIO
from pathlib import Path
from typing import List, Optional

from fastapi import FastAPI, Depends, HTTPException, Query, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from sqlalchemy.orm import Session
from sqlalchemy import text

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Mm, Pt, RGBColor
from lxml import html as lxml_html

from .database import Base, engine, get_db
from .models import (
    Component, FauDb, FcoDb, FcsDb, FdpDb, FiaDb, FmtDb, FprDb, FptDb, FruDb, FtaDb, FtpDb,
    AcoDb, AdvDb, AgdDb, AlcDb, ApeDb, AseDb, AteDb, AvaDb, ElementListDb
)
from .schemas import (
    ComponentCreate, ComponentOut, ComponentUpdate, XmlParseResponse, XmlImportResponse,
    ComponentFamilyOut, ElementListOut
)
from .xml_parser_service import XmlParserService
from pydantic import BaseModel, Field, ConfigDict


@asynccontextmanager
async def lifespan(app: FastAPI):
    # Startup
    Base.metadata.create_all(bind=engine)
    yield
    # Shutdown (if needed)


app = FastAPI(title="CCGenTool2 API", lifespan=lifespan)

COVER_UPLOAD_ROOT = Path(os.getenv("COVER_UPLOAD_DIR", Path(tempfile.gettempdir()) / "ccgentool2_cover_uploads"))
COVER_UPLOAD_ROOT.mkdir(parents=True, exist_ok=True)
COVER_DOCX_ROOT = Path(os.getenv("COVER_DOCX_DIR", Path(tempfile.gettempdir()) / "ccgentool2_cover_docx"))
COVER_DOCX_ROOT.mkdir(parents=True, exist_ok=True)
SFR_DOCX_ROOT = Path(os.getenv("SFR_DOCX_DIR", Path(tempfile.gettempdir()) / "ccgentool2_sfr_docx"))
SFR_DOCX_ROOT.mkdir(parents=True, exist_ok=True)
SAR_DOCX_ROOT = Path(os.getenv("SAR_DOCX_DIR", Path(tempfile.gettempdir()) / "ccgentool2_sar_docx"))
SAR_DOCX_ROOT.mkdir(parents=True, exist_ok=True)
ST_INTRO_DOCX_ROOT = Path(os.getenv("ST_INTRO_DOCX_DIR", Path(tempfile.gettempdir()) / "ccgentool2_stintro_docx"))
ST_INTRO_DOCX_ROOT.mkdir(parents=True, exist_ok=True)
SPD_ASSUMPTIONS_DOCX_ROOT = Path(os.getenv("SPD_ASSUMPTIONS_DOCX_DIR", Path(tempfile.gettempdir()) / "ccgentool2_spd_assumptions_docx"))
SPD_ASSUMPTIONS_DOCX_ROOT.mkdir(parents=True, exist_ok=True)
SPD_THREATS_DOCX_ROOT = Path(os.getenv("SPD_THREATS_DOCX_DIR", Path(tempfile.gettempdir()) / "ccgentool2_spd_threats_docx"))
SPD_THREATS_DOCX_ROOT.mkdir(parents=True, exist_ok=True)
SPD_OSP_DOCX_ROOT = Path(os.getenv("SPD_OSP_DOCX_DIR", Path(tempfile.gettempdir()) / "ccgentool2_spd_osp_docx"))
SPD_OSP_DOCX_ROOT.mkdir(parents=True, exist_ok=True)
FINAL_DOCX_ROOT = Path(os.getenv("FINAL_DOCX_DIR", Path(tempfile.gettempdir()) / "ccgentool2_final_docx"))
FINAL_DOCX_ROOT.mkdir(parents=True, exist_ok=True)

USER_ID_PATTERN = re.compile(r"^[A-Za-z0-9_-]{1,64}$")


def get_user_upload_dir(user_id: str, *, create: bool = False) -> Path:
    if not USER_ID_PATTERN.match(user_id):
        raise HTTPException(status_code=400, detail="Invalid user identifier")

    user_dir = COVER_UPLOAD_ROOT / user_id
    if create:
        user_dir.mkdir(parents=True, exist_ok=True)
    return user_dir


def _get_preview_docx_dir(root: Path, user_id: str, *, create: bool = False) -> Path:
    if not USER_ID_PATTERN.match(user_id):
        raise HTTPException(status_code=400, detail="Invalid user identifier")

    user_dir = root / user_id
    if create:
        user_dir.mkdir(parents=True, exist_ok=True)
    return user_dir


def get_user_docx_dir(user_id: str, *, create: bool = False) -> Path:
    return _get_preview_docx_dir(COVER_DOCX_ROOT, user_id, create=create)


class CoverPreviewRequest(BaseModel):
    model_config = ConfigDict(populate_by_name=True)

    user_id: str = Field(..., alias="user_id")
    title: Optional[str] = None
    version: Optional[str] = None
    revision: Optional[str] = None
    description: Optional[str] = None
    manufacturer: Optional[str] = None
    date: Optional[str] = None
    image_path: Optional[str] = Field(None, alias="image_path")


class HtmlPreviewRequest(BaseModel):
    model_config = ConfigDict(populate_by_name=True)

    user_id: str = Field(..., alias="user_id")
    html_content: str = Field(..., alias="html_content")


class STIntroPreviewRequest(BaseModel):
    model_config = ConfigDict(populate_by_name=True)

    user_id: str = Field(..., alias="user_id")
    cover_data: Optional[dict] = None
    st_reference_html: Optional[str] = None
    toe_reference_html: Optional[str] = None
    toe_overview_html: Optional[str] = None
    toe_description_html: Optional[str] = None


class FinalPreviewRequest(BaseModel):
    model_config = ConfigDict(populate_by_name=True)

    user_id: str = Field(..., alias="user_id")
    cover_data: Optional[dict] = None
    st_reference_html: Optional[str] = None
    toe_reference_html: Optional[str] = None
    toe_overview_html: Optional[str] = None
    toe_description_html: Optional[str] = None
    spd_assumptions_html: Optional[str] = None
    spd_threats_html: Optional[str] = None
    spd_osp_html: Optional[str] = None
    conformance_claims_html: Optional[str] = None
    sfr_list: List[dict] = Field(default_factory=list)
    sar_list: List[dict] = Field(default_factory=list)
    selected_eal: Optional[str] = None
    sfr_preview_html: Optional[str] = None
    sar_preview_html: Optional[str] = None


def _resolve_uploaded_image_path(image_path: Optional[str], user_id: str) -> Optional[Path]:
    if not image_path:
        return None

    expected_prefix = f"/cover/uploads/{user_id}/"
    if not image_path.startswith(expected_prefix):
        raise HTTPException(status_code=400, detail="Invalid image reference for preview generation")

    filename = Path(image_path).name
    upload_dir = get_user_upload_dir(user_id, create=False)
    image_file = upload_dir / filename
    if not image_file.exists():
        raise HTTPException(status_code=400, detail="Referenced cover image could not be found")
    return image_file


def _format_cover_date(date_value: Optional[str]) -> str:
    if not date_value:
        return "—"

    try:
        parsed = datetime.fromisoformat(date_value)
        return parsed.strftime("%d %B %Y")
    except ValueError:
        return date_value


def _build_cover_document(payload: CoverPreviewRequest) -> Path:
    image_file = _resolve_uploaded_image_path(payload.image_path, payload.user_id)
    docx_dir = get_user_docx_dir(payload.user_id, create=True)

    # Clear previous previews for the user to avoid stale documents piling up
    for existing in docx_dir.glob("*.docx"):
        existing.unlink(missing_ok=True)

    document = Document()
    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)

    if image_file:
        image_paragraph = document.add_paragraph()
        image_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = image_paragraph.add_run()
        run.add_picture(str(image_file), width=Mm(120))
        image_paragraph.space_after = Pt(12)

    title_text = payload.title.strip() if payload.title else "Security Target Title"
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_paragraph.add_run(title_text)
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_paragraph.space_after = Pt(12)

    if payload.description:
        description_paragraph = document.add_paragraph(payload.description.strip())
        description_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        description_paragraph.space_after = Pt(18)

    info_items = [
        ("Version", payload.version.strip() if payload.version else "—"),
        ("Revision", payload.revision.strip() if payload.revision else "—"),
        ("Manufacturer/Laboratory", payload.manufacturer.strip() if payload.manufacturer else "—"),
        ("Date", _format_cover_date(payload.date)),
    ]

    for label, value in info_items:
        paragraph = document.add_paragraph()
        paragraph.space_after = Pt(6)
        run_label = paragraph.add_run(f"{label}: ")
        run_label.font.bold = True
        paragraph.add_run(value)

    filename = f"{uuid.uuid4().hex}.docx"
    output_path = docx_dir / filename
    document.save(str(output_path))
    return output_path


def _px_to_points(px_value: float) -> float:
    # Approximate conversion assuming 96px = 72pt
    return px_value * 0.75


def _parse_margin_left(style: Optional[str]) -> Optional[float]:
    if not style:
        return None
    match = re.search(r"margin-left\s*:\s*([0-9.]+)px", style)
    if not match:
        return None
    try:
        return _px_to_points(float(match.group(1)))
    except ValueError:
        return None


def _parse_color(value: Optional[str]) -> Optional[RGBColor]:
    if not value:
        return None

    color = value.strip().lower()
    if color.startswith("#"):
        hex_value = color[1:]
        if len(hex_value) == 3:
            hex_value = "".join(ch * 2 for ch in hex_value)
        if len(hex_value) == 6:
            try:
                r = int(hex_value[0:2], 16)
                g = int(hex_value[2:4], 16)
                b = int(hex_value[4:6], 16)
                return RGBColor(r, g, b)
            except ValueError:
                return None
    elif color.startswith("rgb"):
        numbers = re.findall(r"[0-9]{1,3}", color)
        if len(numbers) >= 3:
            try:
                r, g, b = (min(255, max(0, int(num))) for num in numbers[:3])
                return RGBColor(r, g, b)
            except ValueError:
                return None
    return None


def _collect_inline_styles(element) -> dict:
    styles = {
        "bold": False,
        "italic": False,
        "underline": False,
        "strike": False,
        "color": None,
        "size": None,
    }

    tag = (element.tag or "").lower()
    if tag in {"strong", "b"}:
        styles["bold"] = True
    if tag in {"em", "i"}:
        styles["italic"] = True
    if tag in {"u", "ins"}:
        styles["underline"] = True
    if tag in {"s", "strike", "del"}:
        styles["strike"] = True

    style_attr = element.get("style", "")
    for rule in style_attr.split(";"):
        rule = rule.strip().lower()
        if not rule:
            continue
        if "bold" in rule:
            styles["bold"] = True
        if "italic" in rule:
            styles["italic"] = True
        if "underline" in rule:
            styles["underline"] = True
        if "line-through" in rule:
            styles["strike"] = True
        if rule.startswith("color"):
            parts = rule.split(":", 1)
            if len(parts) == 2:
                parsed = _parse_color(parts[1])
                if parsed:
                    styles["color"] = parsed

    color_attr = element.get("color")
    parsed_color = _parse_color(color_attr)
    if parsed_color:
        styles["color"] = parsed_color

    return styles


CSS_WIDTH_RE = re.compile(r"width\s*:\s*([0-9.]+)px", re.IGNORECASE)
CSS_HEIGHT_RE = re.compile(r"height\s*:\s*([0-9.]+)px", re.IGNORECASE)


def _apply_styles_to_run(run, styles: dict):
    if styles.get("bold"):
        run.bold = True
    if styles.get("italic"):
        run.italic = True
    if styles.get("underline"):
        run.underline = True
    if styles.get("strike"):
        run.strike = True
    color = styles.get("color")
    if color:
        run.font.color.rgb = color
    size = styles.get("size")
    if size:
        run.font.size = size


def _merge_styles(parent: dict, child: dict) -> dict:
    merged = parent.copy()
    for key, value in child.items():
        if key in {"color", "size"}:
            if value is not None:
                merged[key] = value
        elif value:
            merged[key] = True
    return merged


def _px_to_mm(px: float) -> float:
    return px * 0.264583


def _extract_dimension_px(element, attr_name: str) -> Optional[float]:
    if element is None:
        return None

    style = element.get("style", "")
    regex = CSS_WIDTH_RE if attr_name == "width" else CSS_HEIGHT_RE
    match = regex.search(style)
    if match:
        try:
            return float(match.group(1))
        except (TypeError, ValueError):
            pass

    attr_value = element.get(attr_name)
    if attr_value:
        try:
            return float(attr_value)
        except (TypeError, ValueError):
            pass

    if attr_name == "width":
        colwidth = element.get("data-colwidth")
        if colwidth:
            try:
                parts = [float(part) for part in colwidth.split(",") if part.strip()]
                if parts:
                    return parts[0]
            except (TypeError, ValueError):
                pass

    return None


def _decode_base64_image(src: str) -> Optional[bytes]:
    if not src:
        return None

    if src.startswith("data:image"):
        try:
            header, data = src.split(",", 1)
        except ValueError:
            return None
        if ";base64" not in header:
            return None
        try:
            return base64.b64decode(data)
        except Exception:
            return None

    return None


def _append_inline_content(paragraph, element, inherited_styles: Optional[dict] = None):
    styles = _collect_inline_styles(element)
    combined_styles = _merge_styles(inherited_styles or {
        "bold": False,
        "italic": False,
        "underline": False,
        "strike": False,
        "color": None,
        "size": None,
    }, styles)

    tag = (element.tag or "").lower()

    if tag == "img":
        image_data = _decode_base64_image(element.get("src", ""))
        if image_data:
            run = paragraph.add_run()
            _apply_styles_to_run(run, combined_styles)
            image_stream = BytesIO(image_data)
            width_px = _extract_dimension_px(element, "width")
            height_px = _extract_dimension_px(element, "height")
            kwargs = {}
            if width_px:
                kwargs["width"] = Mm(_px_to_mm(width_px))
            if not kwargs and height_px:
                kwargs["height"] = Mm(_px_to_mm(height_px))
            try:
                run.add_picture(image_stream, **kwargs)
            except Exception:
                pass
        tail = element.tail or ""
        if tail:
            run = paragraph.add_run(tail.replace("\xa0", " "))
            _apply_styles_to_run(run, combined_styles)
        return

    text = element.text or ""
    if text:
        run = paragraph.add_run(text.replace("\xa0", " "))
        _apply_styles_to_run(run, combined_styles)

    for child in element:
        _append_inline_content(paragraph, child, combined_styles)
        tail = child.tail or ""
        if tail:
            run = paragraph.add_run(tail.replace("\xa0", " "))
            _apply_styles_to_run(run, combined_styles)


def _append_block_element(document: Document, element, inherited_indent: Optional[float] = None):
    tag = (element.tag or "").lower()
    style_attr = element.get("style")
    margin_left = _parse_margin_left(style_attr)
    indent = margin_left if margin_left is not None else inherited_indent

    heading_map = {
        "h1": Pt(24),
        "h2": Pt(20),
        "h3": Pt(18),
        "h4": Pt(16),
        "h5": Pt(14),
        "h6": Pt(12),
    }

    if tag in heading_map:
        paragraph = document.add_paragraph()
        if indent:
            paragraph.paragraph_format.left_indent = Pt(indent)
        base_styles = {
            "bold": True,
            "italic": False,
            "underline": False,
            "strike": False,
            "color": None,
            "size": heading_map[tag],
        }
        _append_inline_content(paragraph, element, base_styles)
        return

    if tag == "p":
        text_content = (element.text or "") + "".join(
            (child.text or "") + (child.tail or "") for child in element
        )
        if not text_content.strip() and not element.findall("*"):
            document.add_paragraph()
            return

        paragraph = document.add_paragraph()
        if indent:
            paragraph.paragraph_format.left_indent = Pt(indent)
        _append_inline_content(paragraph, element)
        return

    if tag in {"div", "section"}:
        child_indent = indent if indent is not None else inherited_indent
        text = (element.text or "").strip()
        if text:
            paragraph = document.add_paragraph(text)
            if child_indent:
                paragraph.paragraph_format.left_indent = Pt(child_indent)
        for child in element:
            _append_block_element(document, child, child_indent)
        tail = (element.tail or "").strip()
        if tail:
            paragraph = document.add_paragraph(tail)
            if child_indent:
                paragraph.paragraph_format.left_indent = Pt(child_indent)
        return

    if tag in {"ul", "ol"}:
        items = [child for child in element if (child.tag or "").lower() == "li"]
        for idx, child in enumerate(items, start=1):
            paragraph = document.add_paragraph()
            if indent:
                paragraph.paragraph_format.left_indent = Pt(indent)
            prefix = "• " if tag == "ul" else f"{idx}. "
            run = paragraph.add_run(prefix)
            _append_inline_content(paragraph, child, {
                "bold": False,
                "italic": False,
                "underline": False,
                "strike": False,
                "color": None,
                "size": None,
            })
        return

    if tag == "table":
        rows = [row for row in element.findall(".//tr")]
        if not rows:
            return

        max_cells = 0
        table_rows = []
        for row in rows:
            cells = [cell for cell in row if (cell.tag or "").lower() in {"th", "td"}]
            table_rows.append(cells)
            max_cells = max(max_cells, len(cells))

        if max_cells == 0:
            return

        table = document.add_table(rows=len(table_rows), cols=max_cells)
        table.style = "Table Grid"

        for row_index, cells in enumerate(table_rows):
            for col_index, cell in enumerate(cells):
                if col_index >= max_cells:
                    continue
                paragraph = table.cell(row_index, col_index).paragraphs[0]
                paragraph.text = ""
                _append_inline_content(paragraph, cell)
                if (cell.tag or "").lower() == "th":
                    for run in paragraph.runs:
                        run.bold = True

        column_widths = _extract_table_column_widths(element, max_cells)
        for idx, width_px in enumerate(column_widths):
            if width_px is None or idx >= len(table.columns):
                continue
            width_mm = _px_to_mm(width_px)
            column_length = Mm(width_mm)
            table.columns[idx].width = column_length
            for row in table.rows:
                row.cells[idx].width = column_length
        return

    if tag == "br":
        document.add_paragraph()
        return

    # Fallback: treat unknown block elements as paragraphs.
    paragraph = document.add_paragraph()
    if indent:
        paragraph.paragraph_format.left_indent = Pt(indent)
    _append_inline_content(paragraph, element)


def _extract_table_column_widths(table_element, max_cols: int) -> List[Optional[float]]:
    widths: List[Optional[float]] = [None] * max_cols

    colgroup = table_element.find("colgroup")
    if colgroup is not None:
        for idx, col in enumerate(colgroup.findall("col")):
            if idx >= max_cols:
                break
            width_px = _extract_dimension_px(col, "width")
            if width_px:
                widths[idx] = width_px

    for row in table_element.findall(".//tr"):
        cells = [cell for cell in row if (cell.tag or "").lower() in {"th", "td"}]
        for idx, cell in enumerate(cells):
            if idx >= max_cols:
                break
            if widths[idx] is None:
                width_px = _extract_dimension_px(cell, "width")
                if width_px:
                    widths[idx] = width_px

    return widths


def _append_html_to_document(document: Document, html_content: str):
    if not html_content or not html_content.strip():
        return

    try:
        fragment = lxml_html.fragment_fromstring(html_content, create_parent=True)
    except (ValueError, TypeError):
        paragraph = document.add_paragraph(html_content)
        return

    for child in fragment:
        _append_block_element(document, child)


def _build_html_preview_document(html_content: str, user_id: str, root: Path) -> Path:
    docx_dir = _get_preview_docx_dir(root, user_id, create=True)

    for existing in docx_dir.glob("*.docx"):
        existing.unlink(missing_ok=True)

    document = Document()
    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)

    _append_html_to_document(document, html_content)

    filename = f"{uuid.uuid4().hex}.docx"
    output_path = docx_dir / filename
    document.save(str(output_path))
    return output_path


def _build_st_intro_combined_document(payload: STIntroPreviewRequest) -> Path:
    """Build a combined ST Introduction document from all sections."""
    docx_dir = _get_preview_docx_dir(ST_INTRO_DOCX_ROOT, payload.user_id, create=True)

    # Clear previous previews
    for existing in docx_dir.glob("*.docx"):
        existing.unlink(missing_ok=True)

    document = Document()
    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)

    # Add cover page if provided
    if payload.cover_data:
        cover_dict = payload.cover_data
        image_path = cover_dict.get("image_path")
        
        # Add cover image if present
        if image_path:
            try:
                image_file = _resolve_uploaded_image_path(image_path, payload.user_id)
                if image_file:
                    image_paragraph = document.add_paragraph()
                    image_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = image_paragraph.add_run()
                    run.add_picture(str(image_file), width=Mm(120))
                    image_paragraph.space_after = Pt(12)
            except:
                pass  # Skip if image not found
        
        # Add cover title
        title_text = cover_dict.get("title", "").strip() or "Security Target Title"
        title_paragraph = document.add_paragraph()
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run = title_paragraph.add_run(title_text)
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_paragraph.space_after = Pt(12)
        
        # Add cover description
        if cover_dict.get("description"):
            description_paragraph = document.add_paragraph(cover_dict["description"].strip())
            description_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            description_paragraph.space_after = Pt(18)
        
        # Add cover metadata
        info_items = [
            ("Version", cover_dict.get("version", "").strip() or "—"),
            ("Revision", cover_dict.get("revision", "").strip() or "—"),
            ("Manufacturer/Laboratory", cover_dict.get("manufacturer", "").strip() or "—"),
            ("Date", _format_cover_date(cover_dict.get("date"))),
        ]
        
        for label, value in info_items:
            paragraph = document.add_paragraph()
            paragraph.space_after = Pt(6)
            run_label = paragraph.add_run(f"{label}: ")
            run_label.font.bold = True
            paragraph.add_run(value)
        
        # Add page break after cover
        document.add_page_break()

    # Add main heading
    heading = document.add_paragraph()
    heading_run = heading.add_run("1. Security Target Introduction")
    heading_run.font.size = Pt(20)
    heading_run.font.bold = True
    heading.space_after = Pt(12)
    
    # Add introduction text
    intro_text = (
        "This section presents the following information required for a Common Criteria (CC) evaluation:\n"
        "• Identifies the Security Target (ST) and the Target of Evaluation (TOE)\n"
        "• Specifies the security target conventions,\n"
        "• Describes the organization of the security target"
    )
    intro_para = document.add_paragraph(intro_text)
    intro_para.space_after = Pt(12)

    # Add ST Reference section
    if payload.st_reference_html:
        st_ref_heading = document.add_paragraph()
        st_ref_run = st_ref_heading.add_run("1.1 ST Reference")
        st_ref_run.font.size = Pt(18)
        st_ref_run.font.bold = True
        st_ref_heading.space_before = Pt(12)
        st_ref_heading.space_after = Pt(8)
        
        _append_html_to_document(document, payload.st_reference_html)

    # Add TOE Reference section
    if payload.toe_reference_html:
        toe_ref_heading = document.add_paragraph()
        toe_ref_run = toe_ref_heading.add_run("1.2 TOE Reference")
        toe_ref_run.font.size = Pt(18)
        toe_ref_run.font.bold = True
        toe_ref_heading.space_before = Pt(12)
        toe_ref_heading.space_after = Pt(8)
        
        _append_html_to_document(document, payload.toe_reference_html)

    # Add TOE Overview section
    if payload.toe_overview_html:
        toe_overview_heading = document.add_paragraph()
        toe_overview_run = toe_overview_heading.add_run("1.3 TOE Overview")
        toe_overview_run.font.size = Pt(18)
        toe_overview_run.font.bold = True
        toe_overview_heading.space_before = Pt(12)
        toe_overview_heading.space_after = Pt(8)
        
        _append_html_to_document(document, payload.toe_overview_html)

    # Add TOE Description section
    if payload.toe_description_html:
        toe_desc_heading = document.add_paragraph()
        toe_desc_run = toe_desc_heading.add_run("1.4 TOE Description")
        toe_desc_run.font.size = Pt(18)
        toe_desc_run.font.bold = True
        toe_desc_heading.space_before = Pt(12)
        toe_desc_heading.space_after = Pt(8)
        
        _append_html_to_document(document, payload.toe_description_html)

    filename = f"{uuid.uuid4().hex}.docx"
    output_path = docx_dir / filename
    document.save(str(output_path))
    return output_path


def _build_final_combined_document(payload: FinalPreviewRequest) -> Path:
    """Build the complete final Security Target document from all sections."""
    docx_dir = _get_preview_docx_dir(FINAL_DOCX_ROOT, payload.user_id, create=True)

    # Clear previous previews
    for existing in docx_dir.glob("*.docx"):
        existing.unlink(missing_ok=True)

    document = Document()
    section = document.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(25)
    section.right_margin = Mm(25)

    # Page 1: Add cover page if provided
    if payload.cover_data:
        cover_dict = payload.cover_data
        image_path = cover_dict.get("image_path")
        
        # Add cover image if present
        if image_path:
            try:
                image_file = _resolve_uploaded_image_path(image_path, payload.user_id)
                if image_file:
                    image_paragraph = document.add_paragraph()
                    image_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = image_paragraph.add_run()
                    run.add_picture(str(image_file), width=Mm(120))
                    image_paragraph.space_after = Pt(12)
            except:
                pass  # Skip if image not found
        
        # Add cover title
        title_text = cover_dict.get("title", "").strip() or "Security Target Title"
        title_paragraph = document.add_paragraph()
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run = title_paragraph.add_run(title_text)
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_paragraph.space_after = Pt(12)
        
        # Add cover description
        if cover_dict.get("description"):
            description_paragraph = document.add_paragraph(cover_dict["description"].strip())
            description_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            description_paragraph.space_after = Pt(18)
        
        # Add cover metadata
        info_items = [
            ("Version", cover_dict.get("version", "").strip() or "—"),
            ("Revision", cover_dict.get("revision", "").strip() or "—"),
            ("Manufacturer/Laboratory", cover_dict.get("manufacturer", "").strip() or "—"),
            ("Date", _format_cover_date(cover_dict.get("date"))),
        ]
        
        for label, value in info_items:
            paragraph = document.add_paragraph()
            paragraph.space_after = Pt(6)
            run_label = paragraph.add_run(f"{label}: ")
            run_label.font.bold = True
            paragraph.add_run(value)
        
        # Add page break after cover
        document.add_page_break()

    # Page 2: Add ST Introduction heading
    heading = document.add_paragraph()
    heading_run = heading.add_run("1. Security Target Introduction")
    heading_run.font.size = Pt(20)
    heading_run.font.bold = True
    heading.space_after = Pt(12)

    intro_text = (
        "This section presents the following information required for a Common Criteria (CC) evaluation:\n"
        "• Identifies the Security Target (ST) and the Target of Evaluation (TOE)\n"
        "• Specifies the security target conventions,\n"
        "• Describes the organization of the security target"
    )
    intro_para = document.add_paragraph(intro_text)
    intro_para.space_after = Pt(12)

    # Add ST Reference section
    if payload.st_reference_html:
        st_ref_heading = document.add_paragraph()
        st_ref_run = st_ref_heading.add_run("1.1 ST Reference")
        st_ref_run.font.size = Pt(18)
        st_ref_run.font.bold = True
        st_ref_heading.space_before = Pt(12)
        st_ref_heading.space_after = Pt(8)
        _append_html_to_document(document, payload.st_reference_html)

    # Page 3: Add TOE Reference section
    if payload.toe_reference_html:
        toe_ref_heading = document.add_paragraph()
        toe_ref_run = toe_ref_heading.add_run("1.2 TOE Reference")
        toe_ref_run.font.size = Pt(18)
        toe_ref_run.font.bold = True
        toe_ref_heading.space_before = Pt(12)
        toe_ref_heading.space_after = Pt(8)
        _append_html_to_document(document, payload.toe_reference_html)

    # Page 4: Add TOE Overview section
    if payload.toe_overview_html:
        toe_overview_heading = document.add_paragraph()
        toe_overview_run = toe_overview_heading.add_run("1.3 TOE Overview")
        toe_overview_run.font.size = Pt(18)
        toe_overview_run.font.bold = True
        toe_overview_heading.space_before = Pt(12)
        toe_overview_heading.space_after = Pt(8)
        _append_html_to_document(document, payload.toe_overview_html)

    # Page 5: Add TOE Description section
    if payload.toe_description_html:
        toe_desc_heading = document.add_paragraph()
        toe_desc_run = toe_desc_heading.add_run("1.4 TOE Description")
        toe_desc_run.font.size = Pt(18)
        toe_desc_run.font.bold = True
        toe_desc_heading.space_before = Pt(12)
        toe_desc_heading.space_after = Pt(8)
        _append_html_to_document(document, payload.toe_description_html)

    # Section 3: Security Problem Definition
    has_spd_content = (
        payload.spd_assumptions_html or 
        payload.spd_threats_html or 
        payload.spd_osp_html
    )
    if has_spd_content:
        document.add_page_break()
        spd_heading = document.add_paragraph()
        spd_run = spd_heading.add_run("3. Security Problem Definition")
        spd_run.font.size = Pt(20)
        spd_run.font.bold = True
        spd_heading.space_before = Pt(12)
        spd_heading.space_after = Pt(8)
        
        # Add SPD introduction
        intro_para = document.add_paragraph("This chapter identifies the following:")
        intro_para.space_after = Pt(6)
        
        bullet_list = document.add_paragraph(style='List Bullet')
        bullet_list.add_run("Significant assumptions about the TOE's operational environment.")
        bullet_list.space_after = Pt(2)
        
        bullet_list2 = document.add_paragraph(style='List Bullet')
        bullet_list2.add_run("Threats that must be countered by the TOE or its environment")
        bullet_list2.space_after = Pt(6)
        
        notation_para = document.add_paragraph(
            "This document identifies assumptions as A.assumption with \"assumption\" specifying a unique name."
        )
        notation_para.space_after = Pt(2)
        
        notation_para2 = document.add_paragraph(
            "Threats are identified as T.threat with \"threat\" specifying a unique name"
        )
        notation_para2.space_after = Pt(12)
        
        # Add Assumptions
        if payload.spd_assumptions_html:
            _append_html_to_document(document, payload.spd_assumptions_html)
        
        # Add Threats
        if payload.spd_threats_html:
            _append_html_to_document(document, payload.spd_threats_html)
        
        # Add OSP
        if payload.spd_osp_html:
            _append_html_to_document(document, payload.spd_osp_html)

    # Section 4: Conformance Claims
    if payload.conformance_claims_html:
        document.add_page_break()
        conf_heading = document.add_paragraph()
        conf_run = conf_heading.add_run("4. Conformance Claims")
        conf_run.font.size = Pt(20)
        conf_run.font.bold = True
        conf_heading.space_before = Pt(12)
        conf_heading.space_after = Pt(8)
        _append_html_to_document(document, payload.conformance_claims_html)

    # Section 5: Security Requirements
    has_security_requirements = (
        payload.sfr_preview_html or 
        (payload.sfr_list and len(payload.sfr_list) > 0) or
        payload.sar_preview_html or 
        (payload.sar_list and len(payload.sar_list) > 0)
    )
    
    if has_security_requirements:
        document.add_page_break()
        sec_req_heading = document.add_paragraph()
        sec_req_run = sec_req_heading.add_run("5. Security Requirements")
        sec_req_run.font.size = Pt(20)
        sec_req_run.font.bold = True
        sec_req_heading.space_before = Pt(12)
        sec_req_heading.space_after = Pt(12)

    # Section 5.1: Security Functional Requirements
    if payload.sfr_preview_html or (payload.sfr_list and len(payload.sfr_list) > 0):
        sfr_heading = document.add_paragraph()
        sfr_run = sfr_heading.add_run("5.1 Security Functional Requirements")
        sfr_run.font.size = Pt(18)
        sfr_run.font.bold = True
        sfr_heading.space_before = Pt(12)
        sfr_heading.space_after = Pt(12)

        if payload.sfr_preview_html:
            _append_html_to_document(document, payload.sfr_preview_html)
        else:
            for sfr_item in payload.sfr_list:
                if sfr_item.get('preview'):
                    _append_html_to_document(document, sfr_item['preview'])
                    document.add_paragraph().space_after = Pt(12)

    # Section 5.2: Security Assurance Requirements
    if payload.sar_preview_html or (payload.sar_list and len(payload.sar_list) > 0):
        sar_heading = document.add_paragraph()
        sar_run = sar_heading.add_run("5.2 Security Assurance Requirements")
        sar_run.font.size = Pt(18)
        sar_run.font.bold = True
        sar_heading.space_before = Pt(12)
        sar_heading.space_after = Pt(12)

        if payload.sar_preview_html:
            _append_html_to_document(document, payload.sar_preview_html)
        else:
            if payload.selected_eal:
                eal_para = document.add_paragraph()
                eal_run = eal_para.add_run(f"Evaluation Assurance Level: {payload.selected_eal}")
                eal_run.font.bold = True
                eal_para.space_after = Pt(12)

            for sar_item in payload.sar_list:
                if sar_item.get('preview'):
                    _append_html_to_document(document, sar_item['preview'])
                    document.add_paragraph().space_after = Pt(12)

    filename = f"{uuid.uuid4().hex}.docx"
    output_path = docx_dir / filename
    document.save(str(output_path))
    return output_path

# CORS configuration: prefer regex if provided to allow any LAN IP on port 5173
origins = os.getenv("CORS_ORIGINS", "http://localhost:5173,http://127.0.0.1:5173").split(",")
origin_regex = os.getenv(
    "CORS_ORIGIN_REGEX",
    None,
)

cors_kwargs = dict(
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

if origin_regex:
    app.add_middleware(
        CORSMiddleware,
        allow_origin_regex=origin_regex,
        **cors_kwargs,
    )
else:
    app.add_middleware(
        CORSMiddleware,
        allow_origins=origins,
        **cors_kwargs,
    )

app.mount("/cover/uploads", StaticFiles(directory=str(COVER_UPLOAD_ROOT)), name="cover-uploads")
app.mount("/cover/docx", StaticFiles(directory=str(COVER_DOCX_ROOT)), name="cover-docx")
app.mount("/security/sfr/docx", StaticFiles(directory=str(SFR_DOCX_ROOT)), name="sfr-docx")
app.mount("/security/sar/docx", StaticFiles(directory=str(SAR_DOCX_ROOT)), name="sar-docx")
app.mount("/st-intro/docx", StaticFiles(directory=str(ST_INTRO_DOCX_ROOT)), name="st-intro-docx")
app.mount("/spd/assumptions/docx", StaticFiles(directory=str(SPD_ASSUMPTIONS_DOCX_ROOT)), name="spd-assumptions-docx")
app.mount("/spd/threats/docx", StaticFiles(directory=str(SPD_THREATS_DOCX_ROOT)), name="spd-threats-docx")
app.mount("/spd/osp/docx", StaticFiles(directory=str(SPD_OSP_DOCX_ROOT)), name="spd-osp-docx")
app.mount("/final-preview/docx", StaticFiles(directory=str(FINAL_DOCX_ROOT)), name="final-docx")


@app.get("/health")
def health(db: Session = Depends(get_db)):
    started = time.time()
    ok = True
    details = {}
    try:
        db.execute(text("SELECT 1"))
    except Exception as e:
        ok = False
        details["db_error"] = str(e)
    latency_ms = int((time.time() - started) * 1000)
    return {
        "status": "ok" if ok else "degraded",
        "latency_ms": latency_ms,
        "database_url": os.getenv("DATABASE_URL", "unset"),
        "timestamp": int(time.time()),
        "details": details,
    }


# CRUD endpoints
@app.get("/components", response_model=List[ComponentOut])
def list_components(
    q: Optional[str] = Query(None, description="Search across select fields"),
    skip: int = 0,
    limit: int = 100,
    db: Session = Depends(get_db),
):
    query = db.query(Component)
    if q:
        like = f"%{q}%"
        query = query.filter(
            (
                (Component.class_name.ilike(like))
                | (Component.family.ilike(like))
                | (Component.component.ilike(like))
                | (Component.component_name.ilike(like))
                | (Component.element.ilike(like))
                | (Component.element_item.ilike(like))
            )
        )
    return query.offset(skip).limit(limit).all()


@app.post("/components", response_model=ComponentOut, status_code=201)
def create_component(payload: ComponentCreate, db: Session = Depends(get_db)):
    item = Component(
        class_name=payload.class_name,
        family=payload.family,
        component=payload.component,
        component_name=payload.component_name,
        element=payload.element,
        element_item=payload.element_item,
    )
    db.add(item)
    db.commit()
    db.refresh(item)
    return item


@app.get("/components/{item_id}", response_model=ComponentOut)
def get_component(item_id: int, db: Session = Depends(get_db)):
    item = db.get(Component, item_id)
    if not item:
        raise HTTPException(status_code=404, detail="Not found")
    return item


@app.put("/components/{item_id}", response_model=ComponentOut)
def update_component(item_id: int, payload: ComponentUpdate, db: Session = Depends(get_db)):
    item = db.get(Component, item_id)
    if not item:
        raise HTTPException(status_code=404, detail="Not found")
    data = payload.model_dump(exclude_unset=True, by_alias=True)
    if "class" in data:
        item.class_name = data["class"]
        del data["class"]
    for k, v in data.items():
        setattr(item, k, v)
    db.add(item)
    db.commit()
    db.refresh(item)
    return item


@app.delete("/components/{item_id}", status_code=204)
def delete_component(item_id: int, db: Session = Depends(get_db)):
    item = db.get(Component, item_id)
    if not item:
        raise HTTPException(status_code=404, detail="Not found")
    db.delete(item)
    db.commit()
    return None


# XML Parser endpoints
@app.post("/xml/parse", response_model=XmlParseResponse)
async def parse_xml_file(file: UploadFile = File(...)):
    """Parse an uploaded XML file and return the structured data."""
    if not file.filename or not file.filename.lower().endswith('.xml'):
        raise HTTPException(status_code=400, detail="File must be an XML file")
    
    try:
        content = await file.read()
        xml_content = content.decode('utf-8')
        
        parser = XmlParserService()
        result = parser.parse_xml_file(xml_content)
        
        return XmlParseResponse(**result)
    
    except UnicodeDecodeError:
        raise HTTPException(status_code=400, detail="File must be valid UTF-8 encoded XML")
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error parsing XML: {str(e)}")


@app.post("/xml/import", response_model=XmlImportResponse)
async def import_xml_to_database(file: UploadFile = File(...), db: Session = Depends(get_db)):
    """Parse an XML file and import components to the database using family-specific tables."""
    if not file.filename or not file.filename.lower().endswith('.xml'):
        raise HTTPException(status_code=400, detail="File must be an XML file")
    
    try:
        content = await file.read()
        xml_content = content.decode('utf-8')
        
        parser = XmlParserService()
        result = parser.import_to_database(xml_content, db)
        
        return XmlImportResponse(**result)
    
    except UnicodeDecodeError:
        raise HTTPException(status_code=400, detail="File must be valid UTF-8 encoded XML")
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing XML: {str(e)}")


# Family Table Endpoints
@app.get("/families")
def list_family_tables():
    """List all available family tables."""
    tables = {
        "functional": [
            {"name": "fau_db", "description": "Security audit (FAU)"},
            {"name": "fco_db", "description": "Communication (FCO)"},
            {"name": "fcs_db", "description": "Cryptographic support (FCS)"},
            {"name": "fdp_db", "description": "User data protection (FDP)"},
            {"name": "fia_db", "description": "Identification and authentication (FIA)"},
            {"name": "fmt_db", "description": "Security management (FMT)"},
            {"name": "fpr_db", "description": "Privacy (FPR)"},
            {"name": "fpt_db", "description": "Protection of the TSF (FPT)"},
            {"name": "fru_db", "description": "Resource utilisation (FRU)"},
            {"name": "fta_db", "description": "TOE access (FTA)"},
            {"name": "ftp_db", "description": "Trusted path/channels (FTP)"},
        ],
        "assurance": [
            {"name": "aco_db", "description": "Composition (ACO)"},
            {"name": "adv_db", "description": "Development (ADV)"},
            {"name": "agd_db", "description": "Guidance documents (AGD)"},
            {"name": "alc_db", "description": "Life-cycle support (ALC)"},
            {"name": "ape_db", "description": "Protection Profile evaluation (APE)"},
            {"name": "ase_db", "description": "Security Target evaluation (ASE)"},
            {"name": "ate_db", "description": "Tests (ATE)"},
            {"name": "ava_db", "description": "Vulnerability assessment (AVA)"},
        ],
        "special": [
            {"name": "element_list_db", "description": "Element list for colored XML elements"},
        ]
    }
    return tables


def get_family_table_model(table_name: str):
    """Get the SQLAlchemy model for a family table."""
    table_models = {
        "fau_db": FauDb, "fco_db": FcoDb, "fcs_db": FcsDb, "fdp_db": FdpDb,
        "fia_db": FiaDb, "fmt_db": FmtDb, "fpr_db": FprDb, "fpt_db": FptDb,
        "fru_db": FruDb, "fta_db": FtaDb, "ftp_db": FtpDb,
        "aco_db": AcoDb, "adv_db": AdvDb, "agd_db": AgdDb, "alc_db": AlcDb,
        "ape_db": ApeDb, "ase_db": AseDb, "ate_db": AteDb, "ava_db": AvaDb,
        "element_list_db": ElementListDb
    }
    return table_models.get(table_name)


@app.get("/families/{table_name}", response_model=List[ComponentFamilyOut])
def list_family_components(
    table_name: str,
    q: Optional[str] = Query(None, description="Search across select fields"),
    skip: int = 0,
    limit: int = 100,
    db: Session = Depends(get_db),
):
    """List components from a specific family table."""
    model = get_family_table_model(table_name)
    if not model:
        raise HTTPException(status_code=404, detail="Table not found")
    
    query = db.query(model)
    if q and hasattr(model, 'class_field'):
        like = f"%{q}%"
        query = query.filter(
            (
                (model.class_field.ilike(like))
                | (model.family.ilike(like))
                | (model.component.ilike(like))
                | (model.component_name.ilike(like))
                | (model.element.ilike(like))
                | (model.element_item.ilike(like))
            )
        )
    return query.offset(skip).limit(limit).all()


@app.get("/element-lists", response_model=List[ElementListOut])
def list_element_lists(
    q: Optional[str] = Query(None, description="Search across select fields"),
    skip: int = 0,
    limit: int = 100,
    db: Session = Depends(get_db),
):
    """List elements from the element_list_db table."""
    query = db.query(ElementListDb)
    if q:
        like = f"%{q}%"
        query = query.filter(
            (
                (ElementListDb.element.ilike(like))
                | (ElementListDb.element_index.ilike(like))
                | (ElementListDb.item_list.ilike(like))
            )
        )
    return query.offset(skip).limit(limit).all()


@app.get("/element-lists/formatted/{element_id}")
def get_formatted_element_list(element_id: str, db: Session = Depends(get_db)):
    """Get formatted element list for a specific element ID (e.g., fau_gen.1.1)."""
    # Get all element list items for this element
    element_items = db.query(ElementListDb).filter(
        ElementListDb.element == element_id
    ).order_by(ElementListDb.element_index).all()
    
    if not element_items:
        raise HTTPException(status_code=404, detail="Element not found")
    
    # Get the main element text from the component table
    main_element = db.query(FauDb).filter(FauDb.element == element_id).first()
    if not main_element:
        # Try other tables if not found in FauDb
        for model in [FcoDb, FcsDb, FdpDb, FiaDb, FmtDb, FprDb, FptDb, FruDb, FtaDb, FtpDb]:
            main_element = db.query(model).filter(model.element == element_id).first()
            if main_element:
                break
    
    main_text = main_element.element_item if main_element else ""
    
    # Format the response
    formatted_items = []
    for item in element_items:
        formatted_items.append(item.item_list)
    
    return {
        "element": element_id,
        "main_text": main_text,
        "items": formatted_items,
        "formatted_display": f"{element_id} {main_text}\n" + "\n".join(formatted_items)
    }


@app.get("/families/{family_name}/formatted")
def get_formatted_family_elements(
    family_name: str,
    db: Session = Depends(get_db)
):
    """Get formatted element lists for all elements in a family (e.g., fau_gen)."""
    # Get all elements from the family
    family_elements = db.query(ElementListDb).filter(
        ElementListDb.element.like(f"{family_name}%")
    ).all()
    
    if not family_elements:
        raise HTTPException(status_code=404, detail="No elements found for family")
    
    # Group by element
    elements_grouped = {}
    for item in family_elements:
        element_id = item.element
        if element_id not in elements_grouped:
            elements_grouped[element_id] = []
        elements_grouped[element_id].append(item.item_list)
    
    result = []
    for element_id, items in elements_grouped.items():
        # Get main element text
        main_element = db.query(FauDb).filter(FauDb.element == element_id).first()
        if not main_element:
            # Try other tables
            for model in [FcoDb, FcsDb, FdpDb, FiaDb, FmtDb, FprDb, FptDb, FruDb, FtaDb, FtpDb]:
                main_element = db.query(model).filter(model.element == element_id).first()
                if main_element:
                    break
        
        main_text = main_element.element_item if main_element else ""
        
        result.append({
            "element": element_id,
            "main_text": main_text,
            "items": items,
            "formatted_display": f"{element_id} {main_text}\n" + "\n".join(items)
        })
    
    return result


@app.get("/families/{table_name}/count")
def count_family_components(table_name: str, db: Session = Depends(get_db)):
    """Get count of components in a family table."""
    model = get_family_table_model(table_name)
    if not model:
        raise HTTPException(status_code=404, detail="Table not found")
    
    count = db.query(model).count()
    return {"table": table_name, "count": count}


@app.post("/families/{table_name}", response_model=ComponentFamilyOut, status_code=201)
def create_family_component(table_name: str, payload: ComponentCreate, db: Session = Depends(get_db)):
    """Create a new component in a specific family table."""
    model = get_family_table_model(table_name)
    if not model:
        raise HTTPException(status_code=404, detail="Table not found")
    
    # Create item using the family-specific model
    item = model(
        class_field=payload.class_name,
        family=payload.family,
        component=payload.component,
        component_name=payload.component_name,
        element=payload.element,
        element_item=payload.element_item,
    )
    db.add(item)
    db.commit()
    db.refresh(item)
    return item


@app.put("/families/{table_name}/{item_id}", response_model=ComponentFamilyOut)
def update_family_component(table_name: str, item_id: int, payload: ComponentUpdate, db: Session = Depends(get_db)):
    """Update a component in a specific family table."""
    model = get_family_table_model(table_name)
    if not model:
        raise HTTPException(status_code=404, detail="Table not found")
    
    item = db.get(model, item_id)
    if not item:
        raise HTTPException(status_code=404, detail="Not found")
    
    data = payload.model_dump(exclude_unset=True, by_alias=True)
    if "class" in data:
        item.class_field = data["class"]
        del data["class"]
    for k, v in data.items():
        setattr(item, k, v)
    db.add(item)
    db.commit()
    db.refresh(item)
    return item


@app.delete("/families/{table_name}/{item_id}", status_code=204)
def delete_family_component(table_name: str, item_id: int, db: Session = Depends(get_db)):
    """Delete a component from a specific family table."""
    model = get_family_table_model(table_name)
    if not model:
        raise HTTPException(status_code=404, detail="Table not found")

    item = db.get(model, item_id)
    if not item:
        raise HTTPException(status_code=404, detail="Not found")

    db.delete(item)
    db.commit()
    return None


@app.post("/cover/upload")
async def upload_cover_image(
    user_id: str = Query(..., alias="user_id"),
    file: UploadFile = File(...),
):
    """Store a temporary cover image for a user session."""

    user_dir = get_user_upload_dir(user_id, create=True)

    allowed_types = {"image/jpeg", "image/png", "image/bmp", "image/x-ms-bmp"}
    if file.content_type not in allowed_types:
        raise HTTPException(status_code=400, detail="Only jpg, jpeg, png, or bmp files are allowed.")

    # Determine file extension based on mime type to avoid trusting filename blindly
    extension_map = {
        "image/jpeg": ".jpg",
        "image/png": ".png",
        "image/bmp": ".bmp",
        "image/x-ms-bmp": ".bmp",
    }
    suffix = extension_map[file.content_type]

    data = await file.read()
    if not data:
        raise HTTPException(status_code=400, detail="Uploaded file is empty.")

    # Remove any existing files for the session to keep storage minimal
    if user_dir.exists():
        for existing in user_dir.iterdir():
            if existing.is_file():
                existing.unlink()

    filename = f"{uuid.uuid4().hex}{suffix}"
    destination = user_dir / filename
    with destination.open("wb") as buffer:
        buffer.write(data)

    return {"path": f"/cover/uploads/{user_id}/{filename}"}


@app.post("/cover/preview")
async def generate_cover_preview(payload: CoverPreviewRequest):
    if not payload.user_id:
        raise HTTPException(status_code=400, detail="User identifier is required")

    # Ensure the upload directory exists to maintain parity with the image uploads
    get_user_upload_dir(payload.user_id, create=True)

    output_path = _build_cover_document(payload)
    return {"path": f"/cover/docx/{payload.user_id}/{output_path.name}"}


@app.post("/security/sfr/preview")
async def generate_sfr_preview(payload: HtmlPreviewRequest):
    if not payload.user_id:
        raise HTTPException(status_code=400, detail="User identifier is required")

    output_path = _build_html_preview_document(payload.html_content, payload.user_id, SFR_DOCX_ROOT)
    return {"path": f"/security/sfr/docx/{payload.user_id}/{output_path.name}"}


@app.post("/security/sar/preview")
async def generate_sar_preview(payload: HtmlPreviewRequest):
    if not payload.user_id:
        raise HTTPException(status_code=400, detail="User identifier is required")

    output_path = _build_html_preview_document(payload.html_content, payload.user_id, SAR_DOCX_ROOT)
    return {"path": f"/security/sar/docx/{payload.user_id}/{output_path.name}"}


@app.post("/st-intro/preview")
async def generate_st_intro_preview(payload: STIntroPreviewRequest):
    if not payload.user_id:
        raise HTTPException(status_code=400, detail="User identifier is required")

    output_path = _build_st_intro_combined_document(payload)
    return {"path": f"/st-intro/docx/{payload.user_id}/{output_path.name}"}


@app.post("/spd/assumptions/preview")
async def generate_spd_assumptions_preview(payload: HtmlPreviewRequest):
    if not payload.user_id:
        raise HTTPException(status_code=400, detail="User identifier is required")

    output_path = _build_html_preview_document(payload.html_content, payload.user_id, SPD_ASSUMPTIONS_DOCX_ROOT)
    return {"path": f"/spd/assumptions/docx/{payload.user_id}/{output_path.name}"}


@app.post("/spd/threats/preview")
async def generate_spd_threats_preview(payload: HtmlPreviewRequest):
    if not payload.user_id:
        raise HTTPException(status_code=400, detail="User identifier is required")

    output_path = _build_html_preview_document(payload.html_content, payload.user_id, SPD_THREATS_DOCX_ROOT)
    return {"path": f"/spd/threats/docx/{payload.user_id}/{output_path.name}"}


@app.post("/spd/osp/preview")
async def generate_spd_osp_preview(payload: HtmlPreviewRequest):
    if not payload.user_id:
        raise HTTPException(status_code=400, detail="User identifier is required")

    output_path = _build_html_preview_document(payload.html_content, payload.user_id, SPD_OSP_DOCX_ROOT)
    return {"path": f"/spd/osp/docx/{payload.user_id}/{output_path.name}"}


@app.delete("/cover/upload/{user_id}")
async def cleanup_cover_images(user_id: str):
    """Delete all temporary cover images associated with a user session."""

    user_dir = get_user_upload_dir(user_id, create=False)
    if user_dir.exists():
        shutil.rmtree(user_dir)
    docx_dir = get_user_docx_dir(user_id, create=False)
    if docx_dir.exists():
        shutil.rmtree(docx_dir)
    return {"status": "deleted"}


@app.delete("/cover/preview/{user_id}")
async def cleanup_cover_preview(user_id: str):
    """Delete generated cover preview documents for a user session."""

    docx_dir = get_user_docx_dir(user_id, create=False)
    if docx_dir.exists():
        shutil.rmtree(docx_dir)
    return {"status": "deleted"}


@app.delete("/security/sfr/preview/{user_id}")
async def cleanup_sfr_preview(user_id: str):
    docx_dir = _get_preview_docx_dir(SFR_DOCX_ROOT, user_id, create=False)
    if docx_dir.exists():
        shutil.rmtree(docx_dir)
    return {"status": "deleted"}


@app.delete("/security/sar/preview/{user_id}")
async def cleanup_sar_preview(user_id: str):
    docx_dir = _get_preview_docx_dir(SAR_DOCX_ROOT, user_id, create=False)
    if docx_dir.exists():
        shutil.rmtree(docx_dir)
    return {"status": "deleted"}


@app.delete("/st-intro/preview/{user_id}")
async def cleanup_st_intro_preview(user_id: str):
    docx_dir = _get_preview_docx_dir(ST_INTRO_DOCX_ROOT, user_id, create=False)
    if docx_dir.exists():
        shutil.rmtree(docx_dir)
    return {"status": "deleted"}


@app.delete("/spd/assumptions/cleanup/{user_id}/{filename}")
async def cleanup_spd_assumptions_file(user_id: str, filename: str):
    if not USER_ID_PATTERN.match(user_id):
        raise HTTPException(status_code=400, detail="Invalid user identifier")
    
    safe_name = Path(filename).name
    if safe_name != filename:
        raise HTTPException(status_code=400, detail="Invalid file name")
    
    docx_dir = _get_preview_docx_dir(SPD_ASSUMPTIONS_DOCX_ROOT, user_id, create=False)
    file_path = docx_dir / safe_name
    if file_path.exists():
        file_path.unlink()
    return {"status": "deleted"}


@app.get("/spd/assumptions/download/{user_id}/{filename}")
async def download_spd_assumptions_preview(user_id: str, filename: str):
    if not USER_ID_PATTERN.match(user_id):
        raise HTTPException(status_code=400, detail="Invalid user identifier")

    safe_name = Path(filename).name
    if safe_name != filename:
        raise HTTPException(status_code=400, detail="Invalid file name")

    docx_dir = _get_preview_docx_dir(SPD_ASSUMPTIONS_DOCX_ROOT, user_id, create=False)
    file_path = docx_dir / safe_name
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")

    return FileResponse(
        path=str(file_path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="SPD_Assumptions_Preview.docx",
    )


@app.delete("/spd/threats/cleanup/{user_id}/{filename}")
async def cleanup_spd_threats_file(user_id: str, filename: str):
    if not USER_ID_PATTERN.match(user_id):
        raise HTTPException(status_code=400, detail="Invalid user identifier")
    
    safe_name = Path(filename).name
    if safe_name != filename:
        raise HTTPException(status_code=400, detail="Invalid file name")
    
    docx_dir = _get_preview_docx_dir(SPD_THREATS_DOCX_ROOT, user_id, create=False)
    file_path = docx_dir / safe_name
    if file_path.exists():
        file_path.unlink()
    return {"status": "deleted"}


@app.get("/spd/threats/download/{user_id}/{filename}")
async def download_spd_threats_preview(user_id: str, filename: str):
    if not USER_ID_PATTERN.match(user_id):
        raise HTTPException(status_code=400, detail="Invalid user identifier")

    safe_name = Path(filename).name
    if safe_name != filename:
        raise HTTPException(status_code=400, detail="Invalid file name")

    docx_dir = _get_preview_docx_dir(SPD_THREATS_DOCX_ROOT, user_id, create=False)
    file_path = docx_dir / safe_name
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")

    return FileResponse(
        path=str(file_path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="SPD_Threats_Preview.docx",
    )


@app.delete("/spd/osp/cleanup/{user_id}/{filename}")
async def cleanup_spd_osp_file(user_id: str, filename: str):
    if not USER_ID_PATTERN.match(user_id):
        raise HTTPException(status_code=400, detail="Invalid user identifier")
    
    safe_name = Path(filename).name
    if safe_name != filename:
        raise HTTPException(status_code=400, detail="Invalid file name")
    
    docx_dir = _get_preview_docx_dir(SPD_OSP_DOCX_ROOT, user_id, create=False)
    file_path = docx_dir / safe_name
    if file_path.exists():
        file_path.unlink()
    return {"status": "deleted"}


@app.get("/spd/osp/download/{user_id}/{filename}")
async def download_spd_osp_preview(user_id: str, filename: str):
    if not USER_ID_PATTERN.match(user_id):
        raise HTTPException(status_code=400, detail="Invalid user identifier")

    safe_name = Path(filename).name
    if safe_name != filename:
        raise HTTPException(status_code=400, detail="Invalid file name")

    docx_dir = _get_preview_docx_dir(SPD_OSP_DOCX_ROOT, user_id, create=False)
    file_path = docx_dir / safe_name
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")

    return FileResponse(
        path=str(file_path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="SPD_OSP_Preview.docx",
    )


@app.post("/final-preview")
async def generate_final_preview(payload: FinalPreviewRequest):
    if not payload.user_id:
        raise HTTPException(status_code=400, detail="User identifier is required")

    output_path = _build_final_combined_document(payload)
    return {"path": f"/final-preview/docx/{payload.user_id}/{output_path.name}"}


@app.delete("/final-preview/{user_id}")
async def cleanup_final_preview(user_id: str):
    docx_dir = _get_preview_docx_dir(FINAL_DOCX_ROOT, user_id, create=False)
    if docx_dir.exists():
        shutil.rmtree(docx_dir)
    return {"status": "deleted"}


@app.get("/final-preview/download/{user_id}/{filename}")
async def download_final_preview(user_id: str, filename: str):
    if not USER_ID_PATTERN.match(user_id):
        raise HTTPException(status_code=400, detail="Invalid user identifier")

    safe_name = Path(filename).name
    if safe_name != filename:
        raise HTTPException(status_code=400, detail="Invalid file name")

    docx_dir = _get_preview_docx_dir(FINAL_DOCX_ROOT, user_id, create=False)
    file_path = docx_dir / safe_name
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")

    return FileResponse(
        path=str(file_path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="Security_Target_Document.docx",
    )
